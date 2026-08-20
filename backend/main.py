import os
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from docx import Document
from io import BytesIO
import pymupdf
from fastapi.middleware.cors import CORSMiddleware

from chatbot import (
    parse_job_description,
    parse_resume,
    match_resume_to_job_description
)


app = FastAPI()

frontend_url = os.getenv("FRONTEND_URL")
allowed_origins = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "http://localhost:3000",
]
if frontend_url:
    allowed_origins.append(frontend_url.rstrip("/"))

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# --------------------------------------------------
# Extract text from uploaded PDF / DOCX
# --------------------------------------------------

async def extract_text(file: UploadFile | str | None) -> str:
    # No file uploaded
    if file is None:
        return ""

    # If received as text string
    if isinstance(file, str):
        if not file.strip():
            return ""
        try:
            raw_bytes = file.encode('latin1')
            if raw_bytes.startswith(b"%PDF"):
                pdf = pymupdf.open(stream=raw_bytes, filetype="pdf")
                text = "".join(page.get_text() for page in pdf)
                pdf.close()
                return text
            elif raw_bytes.startswith(b"PK\x03\x04"):
                document = Document(BytesIO(raw_bytes))
                return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            pass
        return file

    # Read file as bytes
    file_content = await file.read()

    # Empty file
    if not file_content:
        return ""

    # Get filename
    filename = (file.filename or "").lower()

    # PDF
    if filename.endswith(".pdf") or file_content.startswith(b"%PDF"):
        pdf = pymupdf.open(
            stream=file_content,
            filetype="pdf"
        )
        text = "".join(page.get_text() for page in pdf)
        pdf.close()
        return text

    # DOCX
    elif filename.endswith(".docx") or file_content.startswith(b"PK\x03\x04"):
        document = Document(
            BytesIO(file_content)
        )
        text = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )
        return text

    # Unsupported or plain text
    else:
        try:
            return file_content.decode("utf-8")
        except Exception:
            return ""


# --------------------------------------------------
# Chat endpoint
# --------------------------------------------------

@app.post("/get-questions")
async def get_questions(

    # User's question
    question: str = Form(...),

    # Optional resume (accept UploadFile, str, or None)
    resume: UploadFile | str | None = File(None),

    # Optional job description (accept UploadFile, str, or None)
    job_description: UploadFile | str | None = File(None)

):

    print("\n========== REQUEST START ==========")

    # --------------------------------------------------
    # Check request
    # --------------------------------------------------

    print("1. Endpoint reached")
    print("Question:", question)
    print(
        "Resume:",
        getattr(resume, "filename", type(resume).__name__) if resume else None
    )
    print(
        "Job Description:",
        getattr(job_description, "filename", type(job_description).__name__) if job_description else None
    )

    # --------------------------------------------------
    # Extract resume text
    # --------------------------------------------------

    print("\n2. Extracting resume...")

    resume_text = await extract_text(resume)

    print("Resume extracted successfully")
    print("Resume text length:", len(resume_text))

    # --------------------------------------------------
    # Extract job description text
    # --------------------------------------------------

    print("\n3. Extracting job description...")

    jd_text = await extract_text(job_description)

    print("Job description extracted successfully")
    print("JD text length:", len(jd_text))

    # --------------------------------------------------
    # Parse job description
    # --------------------------------------------------

    print("\n4. Parsing job description...")

    jd = parse_job_description(jd_text)

    print("Job description parsed successfully")

    # --------------------------------------------------
    # Parse resume
    # --------------------------------------------------

    print("\n5. Parsing resume...")

    resume_data = parse_resume(resume_text)

    print("Resume parsed successfully")

    # --------------------------------------------------
    # Create response generator
    # --------------------------------------------------

    print("\n6. Creating response generator...")

    stream_generator = match_resume_to_job_description(
        resume_data,
        jd,
        question
    )

    print("Response generator created successfully")

    # --------------------------------------------------
    # Stream final response to frontend
    # --------------------------------------------------

    print("\n7. Returning StreamingResponse...")

    return StreamingResponse(
        stream_generator,
        media_type="text/plain; charset=utf-8"
    )